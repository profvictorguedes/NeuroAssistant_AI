from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

EXPORT_DIR = Path("exports")
EXPORT_DIR.mkdir(exist_ok=True)


def _apply_arial_12(document: Document) -> None:
    styles = document.styles
    normal_style = styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Arial"
    normal_font.size = Pt(12)
    normal_style.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def _set_run_font(run) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


class WordExportService:
    def export_docx(self, filename: str, input_text: str, output_text: str) -> str:
        if not filename.endswith(".docx"):
            filename = f"{filename}.docx"

        document = Document()
        _apply_arial_12(document)

        heading = document.add_heading("NeuroAssistant AI Session", level=1)
        for run in heading.runs:
            _set_run_font(run)

        h1 = document.add_heading("Input", level=2)
        for run in h1.runs:
            _set_run_font(run)
        p1 = document.add_paragraph(input_text)
        for run in p1.runs:
            _set_run_font(run)

        h2 = document.add_heading("Output", level=2)
        for run in h2.runs:
            _set_run_font(run)
        p2 = document.add_paragraph(output_text)
        for run in p2.runs:
            _set_run_font(run)

        file_path = EXPORT_DIR / filename
        document.save(str(file_path))
        return str(file_path)

    def export_docx_bytes(self, input_text: str, output_text: str) -> bytes:
        document = Document()
        _apply_arial_12(document)

        heading = document.add_heading("NeuroAssistant AI Session", level=1)
        for run in heading.runs:
            _set_run_font(run)

        h1 = document.add_heading("Input", level=2)
        for run in h1.runs:
            _set_run_font(run)
        p1 = document.add_paragraph(input_text)
        for run in p1.runs:
            _set_run_font(run)

        h2 = document.add_heading("Output", level=2)
        for run in h2.runs:
            _set_run_font(run)
        p2 = document.add_paragraph(output_text)
        for run in p2.runs:
            _set_run_font(run)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.read()